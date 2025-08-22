import os
from flask import Flask, render_template, request, redirect, url_for, flash, session
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_cors import CORS
import google.generativeai as genai
from models import db, User, Document, QAHistory
from utils import clean_text
import docx
import fitz  # PyMuPDF

# ----------------------------
# Gemini API Setup
# ----------------------------
genai.configure(api_key="AIzaSyCuyXROjho1GGELH1_5bzUc6bTrcjT-Rpc")
model = genai.GenerativeModel("gemini-1.5-flash")

# ----------------------------
# Flask App Setup
# ----------------------------
app = Flask(__name__)
app.secret_key = "supersecretkey"  # change later for production
CORS(app)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///app.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db.init_app(app)

# ----------------------------
# Database Setup
# ----------------------------
with app.app_context():
    db.create_all()

# ----------------------------
# Routes
# ----------------------------
@app.route("/")
def home():
    return redirect(url_for("dashboard"))

@app.route("/dashboard")
def dashboard():
    if "user_id" not in session:
        return redirect(url_for("login"))

    user = User.query.get(session["user_id"])
    documents = Document.query.filter_by(user_id=user.id).order_by(Document.uploaded_at.desc()).all()

    return render_template("dashboard.html", documents=documents, username=session.get("username"))

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        if User.query.filter_by(username=username).first():
            flash("Username already exists.", "danger")
        else:
            hashed_pw = generate_password_hash(password)
            new_user = User(username=username, password=hashed_pw)
            db.session.add(new_user)
            db.session.commit()
            flash("Registration successful. Please log in.", "success")
            return redirect(url_for("login"))

    return render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password, password):
            session["user_id"] = user.id
            session["username"] = username
            flash("Login successful.", "success")
            return redirect(url_for("dashboard"))
        else:
            flash("Invalid credentials.", "danger")

    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    flash("Logged out successfully.", "info")
    return redirect(url_for("login"))

# Allowed extensions
ALLOWED_EXTENSIONS = {"txt", "pdf", "docx"}

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/upload", methods=["POST"])
def upload():
    if "user_id" not in session:
        flash("Please log in first.", "warning")
        return redirect(url_for("login"))

    file = request.files.get("document")
    if not file or file.filename == "" or not allowed_file(file.filename):
        flash("Invalid or no file selected. Only .txt, .pdf, .docx allowed.", "danger")
        return redirect(url_for("dashboard"))

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    # Extract content based on file type
    ext = filename.rsplit(".", 1)[1].lower()
    content = ""

    try:
        if ext == "txt":
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

        elif ext == "pdf":
            # import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            content = " ".join([page.get_text() for page in doc])


        elif ext == "docx":
            doc = docx.Document(filepath)
            content = " ".join([p.text for p in doc.paragraphs])
    except Exception as e:
        flash(f"Error reading file: {str(e)}", "danger")
        return redirect(url_for("dashboard"))

    # Clean and save content
    content = clean_text(content)

    new_doc = Document(
        filename=filename,
        content=content,
        user_id=session["user_id"]
    )
    db.session.add(new_doc)
    db.session.commit()

    session["document_id"] = new_doc.id
    flash("File uploaded successfully.", "success")
    return redirect(url_for("document_view", filename=filename))

@app.route("/document/<filename>")
def document_view(filename):
    if "user_id" not in session:
        flash("Please log in first.", "warning")
        return redirect(url_for("login"))

    # Get document for this user
    document = Document.query.filter_by(filename=filename, user_id=session["user_id"]).first()
    if not document:
        flash("Document not found.", "danger")
        return redirect(url_for("dashboard"))

    # Store current document in session
    session["document_id"] = document.id

    # Fetch Q&A history for this document (latest first)
    qas = QAHistory.query.filter_by(document_id=document.id).order_by(QAHistory.timestamp.desc()).all()

    return render_template(
        "document.html",
        filename=document.filename,
        qas=qas,
        username=session.get("username")
    )


@app.route("/ask", methods=["POST"])
def ask():
    if "user_id" not in session:
        flash("Please log in first.", "warning")
        return redirect(url_for("login"))

    question = request.form.get("question", "").strip()
    if not question:
        flash("Please enter a question.", "danger")
        return redirect(url_for("dashboard"))

    # Get the current document from session
    document_id = session.get("document_id")
    document = None
    if document_id:
        document = Document.query.filter_by(id=document_id, user_id=session["user_id"]).first()

    # fallback: use the most recent document for this user
    if not document:
        document = Document.query.filter_by(user_id=session["user_id"]).order_by(Document.uploaded_at.desc()).first()
        if document:
            session["document_id"] = document.id

    if not document:
        flash("No document found. Please upload one first.", "danger")
        return redirect(url_for("dashboard"))

    # Improved Prompt for Gemini
    prompt = f"""
You are an AI assistant helping answer questions based on a document.

Document Content:
{document.content}

User Question:
{question}

Instructions:
- If the answer can be found in the document, answer strictly using the document content and mention that it's from the document.
- If not, use your general knowledge to answer helpfully, and mention that it's not found in the document.
- Do not fabricate information.
- Do NOT say the document is irrelevant.
- Keep answers clear and concise.
"""

    response = model.generate_content(prompt)
    answer = response.text if response and response.text else "Sorry, I couldn’t process that."

    # Save Q&A to database
    qa_entry = QAHistory(
        question=question,
        answer=answer,
        document_id=document.id
    )
    db.session.add(qa_entry)
    db.session.commit()

    # Fetch all Q&A for this document only
    qas = QAHistory.query.filter_by(document_id=document.id).order_by(QAHistory.timestamp.desc()).all()

    return render_template(
        "document.html",
        filename=document.filename,
        qas=qas,
        answer=answer,
        username=session.get("username")
    )



# ----------------------------
# Run App
# ----------------------------
if __name__ == "__main__":
    app.run(debug=True)
